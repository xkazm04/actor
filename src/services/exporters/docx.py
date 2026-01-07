"""DOCX exporter for stakeholder sharing."""

from typing import Dict, Any, Optional, List
from datetime import datetime
import json

from .base import BaseExporter, ExportResult


class DOCXExporter(BaseExporter):
    """Export research results to DOCX-compatible format.

    Generates Office Open XML (OOXML) structure as JSON that can be
    converted to DOCX using libraries like python-docx or docx.js.
    Also provides a simple XML representation for direct conversion.
    """

    @property
    def format_name(self) -> str:
        return "docx"

    @property
    def mime_type(self) -> str:
        return "application/json"  # Returns structure for DOCX generation

    @property
    def file_extension(self) -> str:
        return "docx.json"

    def export(
        self,
        research_result: Dict[str, Any],
        title: Optional[str] = None,
        options: Optional[Dict[str, Any]] = None,
    ) -> ExportResult:
        """Export to DOCX-compatible structure.

        The output is a JSON structure that maps to python-docx API calls.
        Use with python-docx:
            from docx import Document
            doc = Document()
            for element in export_result['elements']:
                if element['type'] == 'heading':
                    doc.add_heading(element['text'], level=element['level'])
                elif element['type'] == 'paragraph':
                    doc.add_paragraph(element['text'])
                # ... etc

        Options:
            company_name: Brand name for header
            include_cover_page: Add cover page (default: True)
            include_toc: Add table of contents placeholder (default: True)
        """
        options = options or {}
        branding = self._get_branding(options)

        query = research_result.get("query", "Unknown Query")
        report_title = title or f"Research Report: {query[:60]}"
        template = research_result.get("template", "unknown")
        status = research_result.get("status", "unknown")

        findings = research_result.get("findings", [])
        perspectives = research_result.get("perspectives", [])
        sources = research_result.get("sources", [])
        cost = research_result.get("cost_summary", {})

        include_cover = options.get("include_cover_page", True)
        include_toc = options.get("include_toc", True)

        elements: List[Dict[str, Any]] = []

        # Document properties
        doc_properties = {
            "title": report_title,
            "author": branding["company_name"],
            "subject": query,
            "created": datetime.now().isoformat(),
            "category": f"{template.title()} Research",
        }

        # Cover page
        if include_cover:
            elements.append({
                "type": "cover_page",
                "company_name": branding["company_name"],
                "title": report_title,
                "subtitle": f"{template.title()} Research Analysis",
                "date": self._format_date(),
                "query": query,
            })
            elements.append({"type": "page_break"})

        # Table of Contents placeholder
        if include_toc:
            elements.append({"type": "heading", "text": "Table of Contents", "level": 1})
            elements.append({"type": "toc_field"})
            elements.append({"type": "page_break"})

        # Executive Summary
        elements.append({"type": "heading", "text": "Executive Summary", "level": 1})

        high_conf = [f for f in findings if f.get("confidence_score", 0) >= 0.7]
        summary_text = (
            f"This research investigated \"{query}\" using the {template.title()} "
            f"methodology. The analysis extracted {len(findings)} key findings from "
            f"{len(sources)} sources, with {len(high_conf)} findings rated high confidence. "
            f"Expert analysis was provided from {len(perspectives)} perspectives."
        )
        elements.append({"type": "paragraph", "text": summary_text, "style": "lead"})

        # Summary statistics table
        elements.append({
            "type": "table",
            "style": "summary_table",
            "rows": [
                ["Metric", "Value"],
                ["Total Findings", str(len(findings))],
                ["High Confidence Findings", str(len(high_conf))],
                ["Sources Analyzed", str(len(sources))],
                ["Expert Perspectives", str(len(perspectives))],
                ["Research Template", template.title()],
                ["Status", status.title()],
            ],
        })

        # Top findings in summary
        if findings:
            elements.append({"type": "heading", "text": "Key Takeaways", "level": 2})
            for f in sorted(findings, key=lambda x: x.get("confidence_score", 0), reverse=True)[:5]:
                summary = f.get("summary") or f.get("content", "")[:100]
                conf = f.get("confidence_score", 0.5)
                elements.append({
                    "type": "bullet",
                    "text": f"{summary} (Confidence: {conf:.0%})",
                    "style": "key_finding",
                })

        elements.append({"type": "page_break"})

        # Detailed Findings
        elements.append({"type": "heading", "text": "Detailed Findings", "level": 1})

        # Group by type
        grouped = {}
        for f in findings:
            ft = f.get("finding_type", "other")
            if ft not in grouped:
                grouped[ft] = []
            grouped[ft].append(f)

        for ftype, flist in grouped.items():
            elements.append({
                "type": "heading",
                "text": f"{ftype.replace('_', ' ').title()} ({len(flist)})",
                "level": 2,
            })

            for f in flist:
                summary = f.get("summary") or f.get("content", "")[:60]
                elements.append({"type": "heading", "text": summary, "level": 3})

                content = f.get("content", "")
                elements.append({"type": "paragraph", "text": content})

                conf = f.get("confidence_score", 0.5)
                temporal = f.get("temporal_context", "present")
                elements.append({
                    "type": "paragraph",
                    "text": f"Confidence: {conf:.0%} | Context: {temporal.title()}",
                    "style": "metadata",
                })

        # Expert Analysis
        if perspectives:
            elements.append({"type": "page_break"})
            elements.append({"type": "heading", "text": "Expert Analysis", "level": 1})

            for p in perspectives:
                ptype = p.get("perspective_type", "unknown").replace("_", " ").title()
                elements.append({"type": "heading", "text": f"{ptype} Perspective", "level": 2})

                analysis = p.get("analysis_text", "")
                elements.append({"type": "paragraph", "text": analysis})

                insights = p.get("key_insights", [])
                if insights:
                    elements.append({"type": "heading", "text": "Key Insights", "level": 3})
                    for i in insights:
                        elements.append({"type": "bullet", "text": i})

                warnings = p.get("warnings", [])
                if warnings:
                    elements.append({"type": "heading", "text": "Warnings", "level": 3})
                    for w in warnings:
                        elements.append({"type": "bullet", "text": w, "style": "warning"})

                recs = p.get("recommendations", [])
                if recs:
                    elements.append({"type": "heading", "text": "Recommendations", "level": 3})
                    for r in recs:
                        elements.append({"type": "bullet", "text": r})

        # Sources
        if sources:
            elements.append({"type": "page_break"})
            elements.append({"type": "heading", "text": "Sources", "level": 1})

            sorted_sources = sorted(sources, key=lambda x: x.get("credibility_score", 0), reverse=True)

            source_rows = [["Title", "Domain", "Credibility"]]
            for s in sorted_sources[:30]:
                source_rows.append([
                    s.get("title", "Unknown")[:40],
                    s.get("domain", ""),
                    f"{s.get('credibility_score', 0.5):.0%}",
                ])

            elements.append({"type": "table", "style": "sources_table", "rows": source_rows})

            # Source URLs as endnotes
            elements.append({"type": "heading", "text": "Source URLs", "level": 2})
            for idx, s in enumerate(sorted_sources[:30], 1):
                url = s.get("url", "")
                elements.append({"type": "paragraph", "text": f"[{idx}] {url}", "style": "source_url"})

        # Appendix: Metadata
        elements.append({"type": "page_break"})
        elements.append({"type": "heading", "text": "Appendix: Research Metadata", "level": 1})

        session_id = research_result.get("session_id", "N/A")
        exec_time = research_result.get("execution_time_seconds", 0)
        total_cost = cost.get("total_cost_usd", 0)

        elements.append({
            "type": "table",
            "style": "metadata_table",
            "rows": [
                ["Property", "Value"],
                ["Session ID", session_id],
                ["Query", query],
                ["Template", template.title()],
                ["Status", status.title()],
                ["Execution Time", f"{exec_time:.1f} seconds"],
                ["API Cost", f"${total_cost:.4f}"],
                ["Generated", self._format_date()],
            ],
        })

        # Footer note
        elements.append({
            "type": "paragraph",
            "text": f"This report was generated by {branding['company_name']} on {self._format_date()}. "
                    "The information contained herein is confidential and intended for the recipient only.",
            "style": "footer",
        })

        # Build output structure
        output = {
            "format": "docx",
            "version": "1.0",
            "properties": doc_properties,
            "branding": branding,
            "elements": elements,
            "styles": self._get_style_definitions(branding),
            "conversion_instructions": {
                "python_docx": "Use python-docx library to iterate elements and build document",
                "example": "for el in data['elements']: doc.add_heading(el['text'], el['level']) if el['type']=='heading' else ...",
            },
        }

        return ExportResult(
            format=self.format_name,
            content=json.dumps(output, indent=2, ensure_ascii=False),
            filename=self._generate_filename(report_title, query),
            mime_type=self.mime_type,
            metadata={
                "title": report_title,
                "element_count": len(elements),
                "requires_conversion": True,
            },
        )

    def _get_style_definitions(self, branding: Dict[str, str]) -> Dict[str, Any]:
        """Get style definitions for DOCX generation."""
        return {
            "heading_1": {
                "font_size": 18,
                "font_name": "Calibri",
                "bold": True,
                "color": branding["primary_color"],
                "space_before": 24,
                "space_after": 12,
            },
            "heading_2": {
                "font_size": 14,
                "font_name": "Calibri",
                "bold": True,
                "color": branding["primary_color"],
                "space_before": 18,
                "space_after": 8,
            },
            "heading_3": {
                "font_size": 12,
                "font_name": "Calibri",
                "bold": True,
                "space_before": 12,
                "space_after": 6,
            },
            "paragraph": {
                "font_size": 11,
                "font_name": "Calibri",
                "line_spacing": 1.15,
            },
            "lead": {
                "font_size": 12,
                "font_name": "Calibri",
                "italic": True,
                "color": "#555555",
            },
            "metadata": {
                "font_size": 10,
                "font_name": "Calibri",
                "italic": True,
                "color": "#666666",
            },
            "warning": {
                "font_size": 11,
                "font_name": "Calibri",
                "color": "#b91c1c",
            },
            "source_url": {
                "font_size": 9,
                "font_name": "Consolas",
                "color": branding["accent_color"],
            },
            "footer": {
                "font_size": 9,
                "font_name": "Calibri",
                "italic": True,
                "color": "#888888",
            },
        }
